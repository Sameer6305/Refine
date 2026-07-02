"""
jd_parser.py — Job description → structured requirements parser.

Parses a job description (plain text, .docx, or .pdf) into a ``ParsedJD``
dataclass.  Gemini is called **once per unique JD** (keyed by MD5 hash) and
the result is cached to ``precomputed/parsed_jd.json`` so subsequent runs
are fully offline.

This module is intentionally called at setup / pre-ranking time, *before* the
per-candidate ranking loop, so the single Gemini call is permissible under the
"no API calls during ranking" constraint.

Public API
----------
    parse_jd_from_text(text)   -> ParsedJD
    parse_jd_from_docx(path)   -> ParsedJD
    parse_jd_from_pdf(path)    -> ParsedJD
    cache_parsed_jd(jd, path)  -> None
    load_cached_jd(path)       -> ParsedJD | None
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# ParsedJD dataclass
# ---------------------------------------------------------------------------

@dataclass
class ParsedJD:
    """Structured representation of a job description."""

    raw_text: str
    role_title: str
    required_skills: list[str]
    preferred_skills: list[str]
    disqualifying_signals: list[str]
    min_years_experience: float
    max_years_experience: float
    preferred_locations: list[str]
    notice_period_preference_days: int
    seniority_level: str
    industry_preference: str
    work_mode: str
    role_embedding_text: str
    jd_hash: str

    # Optional rich context fields populated by the second Gemini call
    vibe_signals: list[str] = field(default_factory=list)
    hiring_context: str = ""


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _md5(text: str) -> str:
    """Return the MD5 hex-digest of *text*."""
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def _clean(text: str) -> str:
    """Collapse whitespace and strip leading/trailing blanks."""
    return re.sub(r"\s+", " ", text).strip()


def _safe_list(value: Any, default: list | None = None) -> list:
    """Return *value* as a list, or *default* (empty list) if invalid."""
    if isinstance(value, list):
        return [str(v) for v in value if v]
    return default if default is not None else []


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _safe_str(value: Any, default: str = "") -> str:
    if value is None:
        return default
    return str(value).strip() or default


# ---------------------------------------------------------------------------
# Gemini extraction
# ---------------------------------------------------------------------------

_EXTRACTION_PROMPT = """\
You are an expert technical recruiter and requirements analyst.

Analyse the job description below and return a **single JSON object** with
exactly the keys listed in the schema.  Be thorough but concise — lists
should contain atomic phrases, not full sentences.

### Schema
{
  "role_title": "string — exact job title",
  "required_skills": ["list", "of", "must-have", "technical skills"],
  "preferred_skills": ["nice-to-have skills"],
  "disqualifying_signals": [
    "short phrases describing backgrounds/profiles that should disqualify a candidate",
    "e.g. consulting-only background",
    "CV/speech-only ML experience",
    "no production deployment experience"
  ],
  "min_years_experience": number,
  "max_years_experience": number,
  "preferred_locations": ["City or region names"],
  "notice_period_preference_days": integer (maximum acceptable notice period),
  "seniority_level": "junior | mid | senior | lead | staff | principal",
  "industry_preference": "product_company | consulting | any",
  "work_mode": "remote | hybrid | onsite",
  "vibe_signals": [
    "cultural / working-style signals from the JD",
    "e.g. shipper > researcher",
    "async-first",
    "writing-heavy"
  ]
}

### Rules
- Output **only** valid JSON — no markdown fences, no commentary.
- If a field cannot be determined from the JD, use sensible defaults:
  - lists → []
  - numbers → 0
  - strings → ""
- For `disqualifying_signals`, infer from negative language, emphasis on
  avoiding certain profiles, and domain exclusions in the JD.
- For `notice_period_preference_days`, look for explicit mentions of notice
  period caps or buyout phrases; default to 30 if unspecified.

### Job Description
{jd_text}
"""

_EMBEDDING_PROMPT = """\
You are writing a rich natural-language profile for a job role that will be
used to compute a semantic embedding vector.

Based on the structured requirements below, write a dense 150-200 word
paragraph that captures:
- The core technical stack and required skills
- Preferred background and industry context
- Seniority and experience expectations
- Cultural and working-style expectations

Do NOT include location, salary, or company boilerplate.
Write in third person, present tense.  Output only the paragraph — no
headers, no bullet points.

### Structured Requirements (JSON)
{requirements_json}
"""


def _get_gemini_model():
    """Lazily configure and return the Gemini generative model."""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key or api_key == "your_new_gemini_api_key_here":
        raise RuntimeError(
            "GEMINI_API_KEY is not set. Export a valid key before calling "
            "the JD parser."
        )
    try:
        import google.generativeai as genai  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "google-generativeai is not installed. "
            "Run: pip install google-generativeai"
        ) from exc

    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.5-flash")


def _extract_with_groq(text: str) -> dict:
    """Use Groq (LLaMA 3.3 70B) to extract structured requirements from JD text.

    Groq is free (30 RPM, 500K tokens/day), extremely fast (~1-2 seconds),
    and requires no credit card. Set GROQ_API_KEY to enable.
    See: https://console.groq.com
    """
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set.")
    try:
        from groq import Groq  # type: ignore  # noqa: PLC0415
    except ImportError as exc:
        raise ImportError("groq is not installed. Run: pip install groq") from exc

    client = Groq(api_key=api_key)
    prompt = _EXTRACTION_PROMPT.replace("{jd_text}", text)
    logger.info("Calling Groq for JD extraction (text length=%d chars)", len(text))

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        response_format={"type": "json_object"},
    )
    raw = response.choices[0].message.content.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        logger.warning("Groq returned non-JSON; attempting partial parse. Error: %s", exc)
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
        logger.error("Could not parse Groq response as JSON. Returning empty dict.")
        return {}


def _extract_requirements(text: str) -> dict:
    """Extract JD requirements using the best available LLM.

    Priority: Groq (free, fast) → Gemini → empty dict.
    Caching happens at a higher level (get_or_parse_jd) so this is called
    at most once per unique JD text.
    """
    # 1. Try Groq first — free tier, ~1-2s response time
    groq_key = os.getenv("GROQ_API_KEY")
    if groq_key:
        try:
            return _extract_with_groq(text)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Groq extraction failed (%s); falling back to Gemini.", exc)

    # 2. Fall back to Gemini
    gemini_key = os.getenv("GEMINI_API_KEY")
    if gemini_key and gemini_key not in ("your_new_gemini_api_key_here", "your_gemini_api_key_here"):
        try:
            return _extract_with_gemini(text)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Gemini extraction failed (%s); returning empty dict.", exc)

    logger.warning(
        "No LLM available (set GROQ_API_KEY or GEMINI_API_KEY). "
        "JD will be parsed with keyword fallback only."
    )
    return {}



    """Call Gemini once to extract structured requirements from JD *text*.

    Returns a raw dict matching the schema in ``_EXTRACTION_PROMPT``.
    Handles partial / malformed JSON gracefully.
    """
    model = _get_gemini_model()

    # Use string replace rather than .format() because the prompt template
    # contains literal `{` and `}` JSON braces in its schema example that
    # would otherwise be interpreted as format fields.
    prompt = _EXTRACTION_PROMPT.replace("{jd_text}", text)
    logger.info("Calling Gemini for JD extraction (text length=%d chars)", len(text))

    response = model.generate_content(
        contents=[{"role": "user", "parts": [prompt]}],
        generation_config={"temperature": 0.0, "response_mime_type": "application/json"},
    )

    raw = response.text.strip()

    # Strip any accidental markdown fences
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        logger.warning("Gemini returned non-JSON output; attempting partial parse. Error: %s", exc)
        # Attempt to extract the first {...} block
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
        logger.error("Could not parse Gemini response as JSON. Returning empty dict.")
        return {}


def _generate_embedding_text_with_gemini(requirements: dict) -> str:
    """Optional second Gemini call: generate a rich paragraph for embedding."""
    try:
        model = _get_gemini_model()
        prompt = _EMBEDDING_PROMPT.format(
            requirements_json=json.dumps(requirements, indent=2)
        )
        response = model.generate_content(
            contents=[{"role": "user", "parts": [prompt]}],
            generation_config={"temperature": 0.2},
        )
        return _clean(response.text.strip())
    except Exception as exc:  # noqa: BLE001
        logger.warning("Could not generate embedding text via Gemini: %s", exc)
        return ""


def _build_role_embedding_text(parsed: dict) -> str:
    """Build a dense embedding text from the extracted fields.

    Falls back to a locally-constructed paragraph if the Gemini call fails.
    """
    # Try a rich Gemini-generated description first
    gemini_text = _generate_embedding_text_with_gemini(parsed)
    if gemini_text:
        return gemini_text

    # Fallback: construct a structured paragraph from the parsed fields
    parts: list[str] = []

    role = parsed.get("role_title", "")
    if role:
        parts.append(f"Role: {role}.")

    req = parsed.get("required_skills", [])
    if req:
        parts.append(f"Required skills: {', '.join(req)}.")

    pref = parsed.get("preferred_skills", [])
    if pref:
        parts.append(f"Preferred skills: {', '.join(pref)}.")

    disq = parsed.get("disqualifying_signals", [])
    if disq:
        parts.append(f"Disqualifying signals: {', '.join(disq)}.")

    yoe_min = parsed.get("min_years_experience", 0)
    yoe_max = parsed.get("max_years_experience", 0)
    if yoe_min or yoe_max:
        parts.append(f"Experience: {yoe_min}–{yoe_max} years.")

    seniority = parsed.get("seniority_level", "")
    industry = parsed.get("industry_preference", "")
    if seniority:
        parts.append(f"Seniority: {seniority}.")
    if industry:
        parts.append(f"Industry: {industry}.")

    vibes = parsed.get("vibe_signals", [])
    if vibes:
        parts.append(f"Culture: {', '.join(vibes)}.")

    return _clean(" ".join(parts))


# ---------------------------------------------------------------------------
# ParsedJD construction from Gemini output
# ---------------------------------------------------------------------------

def _build_parsed_jd(raw_text: str, extracted: dict) -> ParsedJD:
    """Construct a ``ParsedJD`` from raw JD text and Gemini-extracted dict.

    All fields default gracefully if absent from *extracted*.
    """
    embedding_text = _build_role_embedding_text(extracted)

    return ParsedJD(
        raw_text=raw_text,
        role_title=_safe_str(extracted.get("role_title"), "Unknown Role"),
        required_skills=_safe_list(extracted.get("required_skills")),
        preferred_skills=_safe_list(extracted.get("preferred_skills")),
        disqualifying_signals=_safe_list(extracted.get("disqualifying_signals")),
        min_years_experience=_safe_float(extracted.get("min_years_experience"), 0.0),
        max_years_experience=_safe_float(extracted.get("max_years_experience"), 0.0),
        preferred_locations=_safe_list(extracted.get("preferred_locations")),
        notice_period_preference_days=_safe_int(
            extracted.get("notice_period_preference_days"), 30
        ),
        seniority_level=_safe_str(extracted.get("seniority_level"), "senior"),
        industry_preference=_safe_str(
            extracted.get("industry_preference"), "product_company"
        ),
        work_mode=_safe_str(extracted.get("work_mode"), "hybrid"),
        role_embedding_text=embedding_text,
        jd_hash=_md5(raw_text),
        vibe_signals=_safe_list(extracted.get("vibe_signals")),
        hiring_context=_safe_str(extracted.get("hiring_context"), ""),
    )


# ---------------------------------------------------------------------------
# Public parse functions
# ---------------------------------------------------------------------------

def parse_jd_from_text(text: str) -> ParsedJD:
    """Parse a plain-text job description into a ``ParsedJD``.

    Calls Gemini once to extract structured requirements, then constructs
    and returns a ``ParsedJD`` dataclass instance.

    Args:
        text: Raw job description as a string.

    Returns:
        ``ParsedJD`` with all fields populated.
    """
    if not text or not text.strip():
        raise ValueError("JD text must not be empty.")

    extracted = _extract_requirements(text)
    return _build_parsed_jd(text, extracted)


def parse_jd_from_docx(path: str) -> ParsedJD:
    """Parse a .docx job description file into a ``ParsedJD``.

    Args:
        path: Filesystem path to the .docx file.

    Returns:
        ``ParsedJD`` with all fields populated.

    Raises:
        FileNotFoundError: If *path* does not exist.
        ImportError: If python-docx is not installed.
    """
    docx_path = Path(path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables (JDs sometimes use table formatting)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    text = "\n".join(paragraphs)
    logger.info("Extracted %d chars from DOCX: %s", len(text), path)
    return parse_jd_from_text(text)


def parse_jd_from_pdf(path: str) -> ParsedJD:
    """Parse a .pdf job description file into a ``ParsedJD``.

    Args:
        path: Filesystem path to the .pdf file.

    Returns:
        ``ParsedJD`` with all fields populated.

    Raises:
        FileNotFoundError: If *path* does not exist.
        ImportError: If pypdf is not installed.
    """
    pdf_path = Path(path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "pypdf is not installed. Run: pip install pypdf"
        ) from exc

    reader = PdfReader(str(pdf_path))
    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)

    text = "\n".join(pages)
    logger.info("Extracted %d chars from PDF: %s", len(text), path)
    return parse_jd_from_text(text)


# ---------------------------------------------------------------------------
# Caching helpers
# ---------------------------------------------------------------------------

def cache_parsed_jd(parsed_jd: ParsedJD, cache_path: str) -> None:
    """Serialise *parsed_jd* to a JSON file at *cache_path*.

    Creates parent directories as needed.

    Args:
        parsed_jd: The ``ParsedJD`` instance to cache.
        cache_path: Destination file path (e.g. ``precomputed/parsed_jd.json``).
    """
    cache_file = Path(cache_path)
    cache_file.parent.mkdir(parents=True, exist_ok=True)

    data = asdict(parsed_jd)
    with cache_file.open("w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2, ensure_ascii=False)

    logger.info("Cached ParsedJD to %s (hash=%s)", cache_path, parsed_jd.jd_hash)


def load_cached_jd(cache_path: str) -> ParsedJD | None:
    """Load a cached ``ParsedJD`` from *cache_path* if it exists.

    Args:
        cache_path: Path to the JSON cache file.

    Returns:
        The deserialised ``ParsedJD``, or ``None`` if the file does not exist
        or is unreadable.
    """
    cache_file = Path(cache_path)
    if not cache_file.exists():
        logger.debug("No cached JD found at %s", cache_path)
        return None

    try:
        with cache_file.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
        parsed = ParsedJD(**data)
        logger.info(
            "Loaded cached ParsedJD from %s (hash=%s)", cache_path, parsed.jd_hash
        )
        return parsed
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to load cached JD from %s: %s", cache_path, exc)
        return None


# ---------------------------------------------------------------------------
# Top-level convenience function (with caching)
# ---------------------------------------------------------------------------

DEFAULT_CACHE_PATH = "precomputed/parsed_jd.json"


def get_or_parse_jd(
    source: str,
    *,
    cache_path: str = DEFAULT_CACHE_PATH,
    force_refresh: bool = False,
) -> ParsedJD:
    """Return a ``ParsedJD`` for *source*, using cache when possible.

    *source* may be:
    - A path to a ``.docx`` file
    - A path to a ``.pdf`` file
    - A path to a ``.txt`` file
    - A raw JD string (if no file extension is detected)

    The result is cached to *cache_path*.  On subsequent calls with the same
    JD content (same MD5 hash), the cached version is returned without
    calling Gemini.

    Args:
        source: Path or raw text of the job description.
        cache_path: Where to store / read the JSON cache.
        force_refresh: If ``True``, ignore any existing cache.

    Returns:
        ``ParsedJD`` instance.
    """
    # Resolve raw text first so we can compute hash for cache validation
    source_path = Path(source)
    if source_path.suffix.lower() in {".docx", ".pdf", ".txt"} and source_path.exists():
        ext = source_path.suffix.lower()
        if ext == ".docx":
            try:
                from docx import Document  # type: ignore
                doc = Document(str(source_path))
                raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                raw_text = source_path.read_text(encoding="utf-8", errors="replace")
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(str(source_path))
                raw_text = "\n".join(
                    (p.extract_text() or "") for p in reader.pages
                )
            except Exception:
                raw_text = source_path.read_text(encoding="utf-8", errors="replace")
        else:
            raw_text = source_path.read_text(encoding="utf-8", errors="replace")
    else:
        # Treat as raw text string
        raw_text = source

    jd_hash = _md5(raw_text)

    # Check cache
    if not force_refresh:
        cached = load_cached_jd(cache_path)
        if cached is not None and cached.jd_hash == jd_hash:
            logger.info(
                "Cache hit for JD hash %s — skipping Gemini call.", jd_hash
            )
            return cached

    # Parse fresh
    logger.info("Cache miss (hash=%s) — calling Gemini.", jd_hash)
    parsed = parse_jd_from_text(raw_text)

    # Persist to cache
    cache_parsed_jd(parsed, cache_path)
    return parsed
